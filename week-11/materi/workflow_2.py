import pdfplumber
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
print("\n=== WORKFLOW 2: PDF TO EXCEL TO WORD ===")
print("\n=== PARSING LAB RESULTS ===")

def parse_lab_results(pdf_path):
    """Extract lab results from PDF"""
    results = {}
    
    with pdfplumber.open(pdf_path) as pdf:
        text = pdf.pages[0].extract_text()
        
        # Parse patient info
        for line in text.split('\n'):
            if "Nama Pasien:" in line:
                results['nama'] = line.split(":")[1].strip()
            elif "No. RM:" in line:
                results['no_rm'] = line.split(":")[1].strip()
            elif "Tanggal:" in line:
                results['tanggal'] = line.split(":")[1].strip()
            elif "Glukosa" in line:
                # Extract numeric value
                parts = line.split(":")
                if len(parts) > 1:
                    value_part = parts[1].strip().split()[0]
                    results['glukosa'] = value_part
            elif "Kolesterol Total" in line:
                parts = line.split(":")
                if len(parts) > 1:
                    value_part = parts[1].strip().split()[0]
                    results['kolesterol'] = value_part
    
    return results

def complete_lab_workflow(pdf_path, excel_path, word_path):
    """Complete workflow: PDF → Excel → Word"""

    # Step 1: Parse PDF
    print("Step 1: Parsing PDF...")
    lab_data = parse_lab_results(pdf_path)

    # Step 2: Save to Excel
    print("Step 2: Saving to Excel...")
    wb = Workbook()
    ws = wb.active
    ws.title = "Lab Results"

    ws.append(["Field", "Value", "Status"])

    # Add data with status
    glukosa = int(lab_data.get('glukosa', 0))
    kolesterol = int(lab_data.get('kolesterol', 0))

    ws.append(["No. RM", lab_data.get('no_rm', ''), ''])
    ws.append(["Nama", lab_data.get('nama', ''), ''])
    ws.append(["Tanggal", lab_data.get('tanggal', ''), ''])
    ws.append(["Glukosa", f"{glukosa} mg/dL",
               "TINGGI" if glukosa >= 126 else "NORMAL"])
    ws.append(["Kolesterol", f"{kolesterol} mg/dL",
               "TINGGI" if kolesterol >= 240 else "NORMAL"])

    wb.save(excel_path)
    print(f"Excel saved: {excel_path}")

    # Step 3: Generate Word report
    print("Step 3: Generating Word report...")
    doc = Document()
    doc.add_heading('INTERPRETASI HASIL LABORATORIUM', 0)

    doc.add_paragraph()
    doc.add_paragraph(f"No. RM: {lab_data.get('no_rm', '')}")
    doc.add_paragraph(f"Nama Pasien: {lab_data.get('nama', '')}")
    doc.add_paragraph(f"Tanggal Pemeriksaan: {lab_data.get('tanggal', '')}")

    doc.add_paragraph()
    doc.add_heading('Hasil Pemeriksaan', level=1)

    # Glukosa interpretation
    p = doc.add_paragraph()
    p.add_run('Glukosa Puasa: ').bold = True
    p.add_run(f"{glukosa} mg/dL ")
    if glukosa >= 126:
        run = p.add_run('(DIABETES)')
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.bold = True
    elif glukosa >= 100:
        run = p.add_run('(PREDIABETES)')
        run.font.color.rgb = RGBColor(255, 165, 0)
        run.bold = True
    else:
        run = p.add_run('(NORMAL)')
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True

    # Kolesterol interpretation
    p = doc.add_paragraph()
    p.add_run('Kolesterol Total: ').bold = True
    p.add_run(f"{kolesterol} mg/dL ")
    if kolesterol >= 240:
        run = p.add_run('(TINGGI)')
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.bold = True
    elif kolesterol >= 200:
        run = p.add_run('(BORDERLINE)')
        run.font.color.rgb = RGBColor(255, 165, 0)
        run.bold = True
    else:
        run = p.add_run('(NORMAL)')
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True

    # Recommendations
    doc.add_paragraph()
    doc.add_heading('Rekomendasi', level=1)

    if glukosa >= 126 or kolesterol >= 240:
        doc.add_paragraph('⚠ PERLU KONSULTASI SEGERA DENGAN DOKTER')
        doc.add_paragraph('• Diet rendah gula dan lemak')
        doc.add_paragraph('• Olahraga teratur minimal 30 menit/hari')
        doc.add_paragraph('• Kontrol rutin setiap 2 minggu')
    else:
        doc.add_paragraph('✓ Hasil dalam batas normal')
        doc.add_paragraph('• Pertahankan pola hidup sehat')
        doc.add_paragraph('• Kontrol rutin 6 bulan')

    doc.save(word_path)
    print(f"Word report saved: {word_path}")

    print("\n✓ Complete workflow finished!")

# Execute complete workflow
complete_lab_workflow(
    "hasil_lab.pdf",
    "hasil_lab_data.xlsx",
    "interpretasi_lab.docx"
)