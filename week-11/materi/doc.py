from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("python-docx berhasil diimport!")

print("=== READING WORD DOCUMENTS ===")

# Create sample document first
doc = Document()
doc.add_heading('Rekam Medis Pasien', 0)

doc.add_heading('Data Pasien', level=1)
doc.add_paragraph('No. RM: RM-2024-001')
doc.add_paragraph('Nama: Ahmad Rizky Pratama')
doc.add_paragraph('Umur: 28 tahun')
doc.add_paragraph('Gender: Laki-laki')

doc.add_heading('Diagnosa', level=1)
doc.add_paragraph('Diabetes Mellitus Type 2')

doc.add_heading('Terapi', level=1)
doc.add_paragraph('1. Metformin 500mg 2x sehari')
doc.add_paragraph('2. Diet rendah gula')
doc.add_paragraph('3. Olahraga teratur 30 menit/hari')

doc.save('rekam_medis.docx')
print("File rekam_medis.docx berhasil dibuat")

# Read the document
print("\n=== MEMBACA DOKUMEN ===")
doc = Document('rekam_medis.docx')

print("Semua paragraphs:")
for i, para in enumerate(doc.paragraphs):
    if para.text:
        print(f"{i+1}. {para.text}")

# Extract specific information
print("\n=== EKSTRAK INFORMASI ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "No. RM:" in text:
        no_rm = text.split(":")[1].strip()
        print(f"Nomor RM: {no_rm}")
    if "Nama:" in text:
        nama = text.split(":")[1].strip()
        print(f"Nama Pasien: {nama}")
    if "Diagnosa" in text and para.style.name.startswith('Heading'):
        if i + 1 < len(doc.paragraphs):
            diagnosa = doc.paragraphs[i + 1].text.strip()
            print(f"Diagnosa: {diagnosa}")
print("\n=== CREATING WORD DOCUMENTS ===")

doc = Document()

# Add title
title = doc.add_heading('SURAT KETERANGAN SAKIT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add content
doc.add_paragraph()  # Empty line
doc.add_paragraph('Yang bertanda tangan di bawah ini:')

# Doctor info
doc.add_paragraph('Nama\t\t: dr. Sari Kusuma, Sp.PD')
doc.add_paragraph('SIP\t\t: 503/SIP/2023')
doc.add_paragraph('Institusi\t: RS Sardjito Yogyakarta')

doc.add_paragraph()  # Empty line
doc.add_paragraph('Menerangkan bahwa:')

# Patient info
doc.add_paragraph('Nama\t\t: Ahmad Rizky Pratama')
doc.add_paragraph('Umur\t\t: 28 tahun')
doc.add_paragraph('Pekerjaan\t: Karyawan Swasta')
doc.add_paragraph('Alamat\t\t: Jl. Kaliurang Km 5, Yogyakarta')

doc.add_paragraph()
doc.add_paragraph(
    'Setelah dilakukan pemeriksaan pada tanggal 16 Januari 2024, '
    'yang bersangkutan dinyatakan sakit dan memerlukan istirahat '
    'selama 3 (tiga) hari terhitung dari tanggal 16-18 Januari 2024.'
)

doc.add_paragraph()
doc.add_paragraph('Demikian surat keterangan ini dibuat untuk dapat dipergunakan sebagaimana mestinya.')

# Signature section
doc.add_paragraph()
doc.add_paragraph('Yogyakarta, 16 Januari 2024')
doc.add_paragraph('Dokter Pemeriksa,')
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('dr. Sari Kusuma, Sp.PD')

doc.save('surat_keterangan_sakit.docx')
print("Surat keterangan sakit berhasil dibuat")

print("\n=== FORMATTING TEXT ===")

doc = Document()

# Title with formatting
title = doc.add_heading('LAPORAN PEMERIKSAAN PASIEN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Formatted paragraphs
doc.add_paragraph()

# Bold text
p = doc.add_paragraph()
p.add_run('Nomor Rekam Medis: ').bold = True
p.add_run('RM-2024-001')

# Different colors
p = doc.add_paragraph()
run = p.add_run('Status: ')
run.bold = True
run = p.add_run('URGENT')
run.font.color.rgb = RGBColor(255, 0, 0)  # Red
run.bold = True

# Different font size
p = doc.add_paragraph()
run = p.add_run('Diagnosis: ')
run.font.size = Pt(12)
run.bold = True
run = p.add_run('Diabetes Mellitus Type 2')
run.font.size = Pt(12)

# Italic text
p = doc.add_paragraph()
run = p.add_run('Catatan: Pasien memerlukan kontrol rutin setiap 2 minggu')
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

doc.save('laporan_formatted.docx')
print("Dokumen dengan formatting berhasil dibuat")

print("\n=== WORKING WITH TABLES ===")

doc = Document()
doc.add_heading('Data Pasien Rawat Jalan', 0)

# Create table
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No. RM'
hdr_cells[1].text = 'Nama'
hdr_cells[2].text = 'Umur'
hdr_cells[3].text = 'Diagnosa'
hdr_cells[4].text = 'Biaya'

# Data rows
patients = [
    ['RM-001', 'Ahmad Rizky', '28', 'Diabetes', 'Rp 750.000'],
    ['RM-002', 'Sari Dewi', '35', 'Hipertensi', 'Rp 450.000'],
    ['RM-003', 'Budi Santoso', '42', 'Gastritis', 'Rp 350.000']
]

for patient in patients:
    row_cells = table.add_row().cells
    for i, value in enumerate(patient):
        row_cells[i].text = value

doc.save('data_pasien_table.docx')
print("Dokumen dengan tabel berhasil dibuat")

print("\n=== TEMPLATE-BASED GENERATION ===")

# Function to replace placeholders
def generate_certificate(template_path, output_path, data):
    doc = Document(template_path)
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in para.text:
                para.text = para.text.replace(f"{{{{{key}}}}}", str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))
    
    doc.save(output_path)
    print(f"Dokumen {output_path} berhasil dibuat")

# Create template first
doc = Document()
doc.add_heading('SURAT RUJUKAN', 0)
doc.add_paragraph()
doc.add_paragraph('Kepada Yth.')
doc.add_paragraph('{{RUMAH_SAKIT_TUJUAN}}')
doc.add_paragraph()
doc.add_paragraph('Dengan hormat,')
doc.add_paragraph()
doc.add_paragraph('Mohon pemeriksaan lebih lanjut untuk pasien:')
doc.add_paragraph('Nama\t\t: {{NAMA_PASIEN}}')
doc.add_paragraph('Umur\t\t: {{UMUR}} tahun')
doc.add_paragraph('Diagnosa\t: {{DIAGNOSA}}')
doc.add_paragraph()
doc.add_paragraph('Terima kasih atas kerjasamanya.')
doc.add_paragraph()
doc.add_paragraph('Hormat kami,')
doc.add_paragraph('{{NAMA_DOKTER}}')

doc.save('template_rujukan.docx')
print("Template rujukan berhasil dibuat")

# Generate documents from template
patients_data = [
    {
        'NAMA_PASIEN': 'Ahmad Rizky',
        'UMUR': 28,
        'DIAGNOSA': 'Diabetes Mellitus Type 2',
        'RUMAH_SAKIT_TUJUAN': 'RS Bethesda',
        'NAMA_DOKTER': 'dr. Sari Kusuma, Sp.PD'
    },
    {
        'NAMA_PASIEN': 'Budi Santoso',
        'UMUR': 42,
        'DIAGNOSA': 'Hipertensi Grade 2',
        'RUMAH_SAKIT_TUJUAN': 'RS Panti Rapih',
        'NAMA_DOKTER': 'dr. Sari Kusuma, Sp.PD'
    }
]

for i, patient_data in enumerate(patients_data):
    generate_certificate(
        'template_rujukan.docx',
        f'surat_rujukan_{i+1}.docx',
        patient_data
    )

print(f"\nTotal {len(patients_data)} surat rujukan berhasil dibuat!")
