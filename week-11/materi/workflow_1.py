print("\n=== WORKFLOW 1: EXCEL TO WORD REPORT ===")

from openpyxl import load_workbook
from docx import Document
from datetime import datetime

def generate_high_risk_report(excel_path, output_docx):
    """Generate Word report for high-risk patients from Excel"""

    # Read Excel
    wb = load_workbook(excel_path)
    ws = wb.active

    # Create Word document
    doc = Document()
    doc.add_heading('LAPORAN PASIEN RISIKO TINGGI', 0)

    # Add date
    doc.add_paragraph(f'Tanggal: {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph()

    # Filter and collect high-risk patients
    high_risk_patients = []

    for row in range(2, ws.max_row + 1):
        umur = ws.cell(row=row, column=3).value
        diagnosa = ws.cell(row=row, column=5).value
        biaya = ws.cell(row=row, column=6).value

        # High risk criteria: age > 35 OR expensive treatment
        if umur > 35 or biaya > 600000:
            patient = {
                'no_rm': ws.cell(row=row, column=1).value,
                'nama': ws.cell(row=row, column=2).value,
                'umur': umur,
                'gender': ws.cell(row=row, column=4).value,
                'diagnosa': diagnosa,
                'biaya': biaya
            }
            high_risk_patients.append(patient)

    # Add summary
    doc.add_heading('Ringkasan', level=1)
    doc.add_paragraph(f'Total pasien risiko tinggi: {len(high_risk_patients)} orang')
    doc.add_paragraph()

    # Add table
    doc.add_heading('Daftar Pasien', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No. RM'
    hdr_cells[1].text = 'Nama'
    hdr_cells[2].text = 'Umur'
    hdr_cells[3].text = 'Diagnosa'
    hdr_cells[4].text = 'Biaya'

    # Add patients
    for patient in high_risk_patients:
        row_cells = table.add_row().cells
        row_cells[0].text = patient['no_rm']
        row_cells[1].text = patient['nama']
        row_cells[2].text = str(patient['umur'])
        row_cells[3].text = patient['diagnosa']
        row_cells[4].text = f"Rp {patient['biaya']:,}"

    # Add recommendations
    doc.add_paragraph()
    doc.add_heading('Rekomendasi', level=1)
    doc.add_paragraph('1. Follow-up intensif untuk pasien di atas 40 tahun')
    doc.add_paragraph('2. Monitoring biaya treatment untuk kasus kompleks')
    doc.add_paragraph('3. Case conference untuk pasien dengan biaya > Rp 700.000')

    # Save
    doc.save(output_docx)
    print(f"Report saved: {output_docx}")
    print(f"Total high-risk patients: {len(high_risk_patients)}")

    return len(high_risk_patients)

# Generate report
count = generate_high_risk_report("data_pasien.xlsx", "laporan_risiko_tinggi.docx")