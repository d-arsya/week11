import pdfplumber
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from docx import Document
from datetime import datetime
print("\n=== WORKFLOW 3: BULK CERTIFICATE GENERATION ===")

def bulk_generate_certificates(excel_path, template_path, output_folder):
    """Generate certificates for all patients in Excel"""

    # Create output folder if not exists
    import os
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Read Excel
    wb = load_workbook(excel_path)
    ws = wb.active

    # Counter
    count = 0

    # Process each patient
    for row in range(2, ws.max_row + 1):
        patient_data = {
            'NAMA_PASIEN': ws.cell(row=row, column=2).value,
            'UMUR': ws.cell(row=row, column=3).value,
            'NO_RM': ws.cell(row=row, column=1).value,
            'DIAGNOSA': ws.cell(row=row, column=5).value,
            'TANGGAL': datetime.now().strftime("%d %B %Y"),
            'DOKTER': 'dr. Sari Kusuma, Sp.PD'
        }

        # Create certificate
        doc = Document()
        doc.add_heading('SURAT KETERANGAN SAKIT', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph('Yang bertanda tangan di bawah ini:')
        doc.add_paragraph(f'Nama\t\t: {patient_data["DOKTER"]}')
        doc.add_paragraph('Institusi\t: RS Sardjito Yogyakarta')

        doc.add_paragraph()
        doc.add_paragraph('Menerangkan bahwa:')
        doc.add_paragraph(f'Nama\t\t: {patient_data["NAMA_PASIEN"]}')
        doc.add_paragraph(f'Umur\t\t: {patient_data["UMUR"]} tahun')
        doc.add_paragraph(f'No. RM\t\t: {patient_data["NO_RM"]}')

        doc.add_paragraph()
        doc.add_paragraph(
            f'Pasien tersebut didiagnosa {patient_data["DIAGNOSA"]} '
            f'dan memerlukan istirahat selama 3 hari.'
        )

        doc.add_paragraph()
        doc.add_paragraph(f'{patient_data["TANGGAL"]}')
        doc.add_paragraph('Dokter Pemeriksa,')
        doc.add_paragraph()
        doc.add_paragraph(patient_data["DOKTER"])

        # Save with patient name
        output_file = os.path.join(
            output_folder,
            f"Surat_{patient_data['NO_RM']}_{patient_data['NAMA_PASIEN'].replace(' ', '_')}.docx"
        )
        doc.save(output_file)

        count += 1
        print(f"Generated: {output_file}")

    print(f"\n✓ Total {count} certificates generated in folder: {output_folder}")
    return count

# Generate bulk certificates
total = bulk_generate_certificates(
    "data_pasien.xlsx",
    "template_surat.docx",
    "surat_keterangan"
)